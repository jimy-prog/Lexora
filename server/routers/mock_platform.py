from fastapi import APIRouter, Request, Depends, HTTPException, Form, UploadFile, File
import shutil
import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from datetime import datetime
import json
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from master_database import SessionMaster, MockExam, ExamSection, QuestionBlock, Question, AnswerOption, MockAttempt, AttemptAnswer, WritingFeedback, ClassRoom, ClassMember, User
from auth import get_current_user
from config import UPLOADS_DIR, GEMINI_API_KEY

router = APIRouter(prefix="/mock", tags=["mock"])
templates = Jinja2Templates(directory="templates")

# Lazy-load Gemini to avoid startup crash if key is missing
def _get_gemini_model():
    try:
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception:
        return None

def get_mdb():
    db = SessionMaster()
    try:
        yield db
    finally:
        db.close()

@router.get("/", response_class=HTMLResponse)
async def mock_dashboard(request: Request, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login?next=/mock", status_code=302)
        
    if user.role == "owner":
        exams = db.query(MockExam).all()
    else:
        exams = db.query(MockExam).filter(MockExam.is_published == True).all()
    
    return templates.TemplateResponse("mock_dashboard.html", {
        "request": request,
        "active_page": "mock",
        "user": user,
        "exams": exams
    })

@router.get("/manage", response_class=HTMLResponse)
async def manage_mocks(request: Request, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=302)
        
    if user.role != "owner":
        return RedirectResponse("/mock/", status_code=302)
        
    exams = db.query(MockExam).order_by(MockExam.id.desc()).all()
    
    return templates.TemplateResponse("manage_mocks.html", {
        "request": request,
        "active_page": "manage_mocks",
        "user": user,
        "exams": exams
    })

@router.post("/create")
async def create_mock(request: Request, title: str = Form(...), exam_type: str = Form("IELTS"), test_scope: str = Form("Full Test"), test_mode: str = Form("Exam Mode"), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Only Admins can create exams")
        
    exam = MockExam(title=title, exam_type=exam_type, test_scope=test_scope, test_mode=test_mode)
    db.add(exam)
    db.commit()
    db.refresh(exam)
    
    return RedirectResponse(f"/mock/{exam.id}/build", status_code=302)

@router.post("/{exam_id}/update-settings")
async def update_settings(request: Request, exam_id: int, title: str = Form(...), test_scope: str = Form(...), test_mode: str = Form(...), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    exam.title = title
    exam.test_scope = test_scope
    exam.test_mode = test_mode
    db.commit()
    
    return RedirectResponse(f"/mock/{exam.id}/build", status_code=303)

# (add-section is defined below with the full implementation)

@router.post("/{exam_id}/add-question")
async def add_question(request: Request, exam_id: int, section_id: int = Form(...), instructions: str = Form(""), q_type: str = Form(...), prompt: str = Form(...), correct_answer_text: str = Form(""), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    section = db.query(ExamSection).filter(ExamSection.id == section_id, ExamSection.exam_id == exam_id).first()
    if not section:
        raise HTTPException(status_code=404, detail="Section not found")
        
    # Get or create a block to hold this question
    # We'll just attach it to the last block in the section, or create one if none exists
    block = db.query(QuestionBlock).filter(QuestionBlock.section_id == section.id).order_by(QuestionBlock.id.desc()).first()
    if not block:
        block = QuestionBlock(section_id=section.id, passage_text="", instructions=instructions)
        db.add(block)
        db.flush()
    elif instructions and instructions != block.instructions:
        # If they provided new instructions, create a new block
        block = QuestionBlock(section_id=section.id, passage_text="", instructions=instructions)
        db.add(block)
        db.flush()
        
    question_num = db.query(Question).join(QuestionBlock).join(ExamSection).filter(ExamSection.exam_id == exam_id).count() + 1
    
    question = Question(
        block_id=block.id,
        q_type=q_type,
        question_number=int(question_num),
        prompt=prompt,
        correct_answer_text=correct_answer_text,
        points=1
    )
    db.add(question)
    db.commit()
    
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.get("/{exam_id}/build", response_class=HTMLResponse)
async def mock_builder(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Only Admins can build exams")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    return templates.TemplateResponse("mock_builder.html", {
        "request": request,
        "active_page": "mock",
        "user": user,
        "exam": exam
    })

@router.post("/upload-speaking")
async def upload_speaking_audio(
    attempt_id: int = Form(...),
    question_id: int = Form(...),
    audio_file: UploadFile = File(...),
    db: SessionMaster = Depends(get_mdb)
):
    # Create directory for attempts if not exists
    attempt_dir = os.path.join(UPLOADS_DIR, f"attempt_{attempt_id}")
    os.makedirs(attempt_dir, exist_ok=True)
    
    file_path = os.path.join(attempt_dir, f"q_{question_id}_{audio_file.filename}")
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(audio_file.file, buffer)
        
    # Update database
    # We find or create the answer for this question
    ans = db.query(AttemptAnswer).filter(
        AttemptAnswer.attempt_id == attempt_id,
        AttemptAnswer.question_id == question_id
    ).first()
    
    if not ans:
        ans = AttemptAnswer(attempt_id=attempt_id, question_id=question_id)
        db.add(ans)
        
    ans.audio_url = f"/uploads/attempt_{attempt_id}/q_{question_id}_{audio_file.filename}"
    db.commit()
    
    return {"status": "success", "url": ans.audio_url}

@router.post("/upload-writing-image")
async def upload_writing_image(
    attempt_id: int = Form(...),
    section_id: int = Form(...),
    image_file: UploadFile = File(...),
    db: SessionMaster = Depends(get_mdb)
):
    attempt_dir = os.path.join(UPLOADS_DIR, f"attempt_{attempt_id}")
    os.makedirs(attempt_dir, exist_ok=True)
    
    filename = f"writing_{section_id}_{image_file.filename}"
    file_path = os.path.join(attempt_dir, filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(image_file.file, buffer)
        
    # Find section to get first question
    section = db.query(ExamSection).filter(ExamSection.id == section_id).first()
    if not section or not section.blocks or not section.blocks[0].questions:
        raise HTTPException(status_code=400, detail="No question found in section to attach upload to")
        
    first_q = section.blocks[0].questions[0]
    
    ans = db.query(AttemptAnswer).filter(
        AttemptAnswer.attempt_id == attempt_id,
        AttemptAnswer.question_id == first_q.id
    ).first()
    
    if not ans:
        ans = AttemptAnswer(attempt_id=attempt_id, question_id=first_q.id)
        db.add(ans)
        
    ans.file_url = f"/uploads/attempt_{attempt_id}/{filename}"
    db.commit()
    
    return {"status": "success", "url": ans.file_url}

@router.post("/{exam_id}/publish")
async def publish_exam(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    exam.is_published = True
    db.commit()
    
    return RedirectResponse("/mock/manage", status_code=303)

@router.post("/{exam_id}/delete")
async def delete_exam(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    db.delete(exam)
    db.commit()
    
    return RedirectResponse("/mock/manage", status_code=303)

@router.get("/{exam_id}/start", response_class=HTMLResponse)
async def take_mock_start(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=302)
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    return templates.TemplateResponse("mock_start.html", {
        "request": request,
        "user": user,
        "exam": exam,
        "active_page": "mock"
    })

@router.get("/{exam_id}/take", response_class=HTMLResponse)
async def mock_take(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(f"/login?next=/mock/{exam_id}/take", status_code=302)
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    # Create or get existing in-progress attempt
    attempt = db.query(MockAttempt).filter(
        MockAttempt.student_id == user.id,
        MockAttempt.exam_id == exam_id,
        MockAttempt.status == "in_progress"
    ).first()
    
    if not attempt:
        attempt = MockAttempt(
            tenant_id=user.tenant_id,
            student_id=user.id,
            exam_id=exam_id,
            status="in_progress"
        )
        db.add(attempt)
        db.commit()
        db.refresh(attempt)
        
    # Fetch teachers
    from master_database import User as MasterUser
    teachers = db.query(MasterUser).filter(MasterUser.account_type == "teacher").all()
        
    return templates.TemplateResponse("mock_take.html", {
        "request": request,
        "user": user,
        "exam": exam,
        "attempt": attempt,
        "teachers": teachers
    })


from fastapi import BackgroundTasks
from fastapi.responses import JSONResponse

# Global dict to track background extraction status
EXTRACTION_STATUS = {}

def process_pdf_background(temp_path: str, exam_id: int, test_scope: str):
    def update_progress(percent: int, message: str):
        EXTRACTION_STATUS[exam_id] = {"progress": percent, "message": message, "status": "processing"}
        
    try:
        update_progress(5, "Initializing AI pipeline...")
        result = extract_ielts_exam_from_pdf(temp_path, test_scope=test_scope, progress_callback=update_progress)
        
        if "error" in result and result["sections"] == []:
            EXTRACTION_STATUS[exam_id] = {"progress": 100, "message": f"AI Error: {result['error']}", "status": "error"}
            return

        update_progress(95, "Building internal database structures...")
        db = SessionMaster()
        try:
            if "sections" in result:
                for s_data in result["sections"]:
                    section = ExamSection(
                        exam_id=exam_id,
                        section_type=s_data.get("section_type", "Reading Section"),
                        time_limit_minutes=20
                    )
                    db.add(section)
                    db.flush()
                    
                    for b_data in s_data.get("blocks", []):
                        block = QuestionBlock(
                            section_id=section.id,
                            instructions=b_data.get("instructions", ""),
                            passage_text=b_data.get("passage_text", "")
                        )
                        db.add(block)
                        db.flush()
                        
                        for idx, q_data in enumerate(b_data.get("questions", [])):
                            question = Question(
                                block_id=block.id,
                                q_type=q_data.get("q_type", "GAP_FILL"),
                                question_number=q_data.get("question_number", idx+1),
                                prompt=q_data.get("prompt", ""),
                                correct_answer_text=q_data.get("correct_answer_text", "")
                            )
                            db.add(question)
                            db.flush()
                            
                            for o_data in q_data.get("options", []):
                                opt = AnswerOption(
                                    question_id=question.id,
                                    text=o_data.get("text", ""),
                                    is_correct=o_data.get("is_correct", False)
                                )
                                db.add(opt)
                
                db.commit()
            EXTRACTION_STATUS[exam_id] = {"progress": 100, "message": "Extraction complete!", "status": "completed"}
        except Exception as e:
            EXTRACTION_STATUS[exam_id] = {"progress": 100, "message": f"Database error: {str(e)}", "status": "error"}
            db.rollback()
        finally:
            db.close()
    except Exception as e:
        EXTRACTION_STATUS[exam_id] = {"progress": 100, "message": f"Pipeline crash: {str(e)}", "status": "error"}

@router.post("/{exam_id}/upload-pdf")
async def process_pdf(request: Request, exam_id: int, background_tasks: BackgroundTasks, pdf_doc: UploadFile = File(...), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Only Admins can process tests")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    os.makedirs("uploads/tmp", exist_ok=True)
    temp_path = f"uploads/tmp/{pdf_doc.filename}"
    with open(temp_path, "wb") as buffer:
        import shutil
        shutil.copyfileobj(pdf_doc.file, buffer)
        
    # Start tracking
    EXTRACTION_STATUS[exam.id] = {"progress": 0, "message": "Queuing document...", "status": "processing"}
    background_tasks.add_task(process_pdf_background, temp_path, exam.id, exam.test_scope)
    
    return JSONResponse({"status": "started"})

@router.get("/{exam_id}/upload-status")
async def get_upload_status(exam_id: int, request: Request):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    status = EXTRACTION_STATUS.get(exam_id, {"progress": 0, "message": "Waiting...", "status": "idle"})
    return JSONResponse(status)

@router.post("/{exam_id}/upload-audio")
async def upload_audio(request: Request, exam_id: int, audio_file: UploadFile = File(...), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Only Admins can upload audio")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    target_dir = UPLOADS_DIR / "audio"
    os.makedirs(target_dir, exist_ok=True)
    filename = f"{exam_id}_{audio_file.filename}"
    file_path = target_dir / filename
    with open(file_path, "wb") as buffer:
        import shutil
        shutil.copyfileobj(audio_file.file, buffer)
        
    exam.audio_url = f"/uploads/audio/{filename}"
    db.commit()
    
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.post("/{exam_id}/upload-image")
async def upload_image(request: Request, exam_id: int, image_file: UploadFile = File(...), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Only Admins can upload images")
        
    target_dir = UPLOADS_DIR / "images"
    os.makedirs(target_dir, exist_ok=True)
    import time
    filename = f"{int(time.time())}_{image_file.filename}"
    file_path = target_dir / filename
    with open(file_path, "wb") as buffer:
        import shutil
        shutil.copyfileobj(image_file.file, buffer)
        
    return RedirectResponse(f"/mock/{exam_id}/build?uploaded_img=/uploads/images/{filename}", status_code=303)

@router.post("/{exam_id}/save-answers")
async def save_answers(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    form = await request.form()
    
    # Process inputs like answer_14, option_23_4
    for key, val in form.items():
        if key.startswith("answer_"):
            q_id = int(key.split("_")[1])
            question = db.query(Question).filter(Question.id == q_id).first()
            if question:
                question.correct_answer_text = val
        elif key.startswith("option_"):
            # expects option_{q_id}_{opt_id}
            _, q_id, opt_id = key.split("_")
            q_id, opt_id = int(q_id), int(opt_id)
            option = db.query(AnswerOption).filter(AnswerOption.id == opt_id).first()
            if option:
                # set all other options for this q to false first, but only if val is "on"
                if val == "on":
                    db.query(AnswerOption).filter(AnswerOption.question_id == q_id).update({"is_correct": False})
                    option.is_correct = True
                    
    db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.get("/question/{q_id}")
async def get_question(q_id: int, db: SessionMaster = Depends(get_mdb)):
    q = db.query(Question).filter(Question.id == q_id).first()
    if not q:
        return JSONResponse({"error": "Not found"}, status_code=404)
    opts = [{"id": o.id, "text": o.text, "is_correct": o.is_correct} for o in q.options]
    return JSONResponse({
        "id": q.id,
        "q_type": q.q_type,
        "prompt": q.prompt,
        "correct_answer_text": q.correct_answer_text,
        "options": opts
    })

@router.post("/{exam_id}/update-question/{q_id}")
async def update_question(request: Request, exam_id: int, q_id: int, 
                          q_type: str = Form(...), prompt: str = Form(...), 
                          correct_answer_text: str = Form(""), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    question = db.query(Question).filter(Question.id == q_id).first()
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
        
    question.q_type = q_type
    question.prompt = prompt
    question.correct_answer_text = correct_answer_text
    
    # Handle MCQ options if provided in form
    form = await request.form()
    options_json = form.get("options_json")
    if options_json:
        try:
            opts = json.loads(options_json)
            # Clear existing
            db.query(AnswerOption).filter(AnswerOption.question_id == q_id).delete()
            for o in opts:
                db.add(AnswerOption(
                    question_id=q_id,
                    text=o['text'],
                    is_correct=o['is_correct'],
                    order=o.get('order', 0)
                ))
        except: pass
        
    db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.post("/{exam_id}/add-section")
async def add_section(request: Request, exam_id: int, 
                      section_type: str = Form(...), time_limit: int = Form(20), 
                      db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
        
    # Find max order
    max_order = db.query(ExamSection).filter(ExamSection.exam_id == exam_id).count()
    
    section = ExamSection(
        exam_id=exam_id,
        section_type=section_type,
        time_limit_minutes=time_limit,
        order=max_order + 1
    )
    db.add(section)
    db.flush()
    
    # Create an empty block for it
    block = QuestionBlock(
        section_id=section.id,
        instructions="Complete the tasks below",
        passage_text="<p>Enter your passage or content here.</p>"
    )
    db.add(block)
    db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

from pydantic import BaseModel
class AIChatRequest(BaseModel):
    message: str

@router.post("/{exam_id}/ai-chat")
async def ai_chat(request: Request, exam_id: int, chat_req: AIChatRequest, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        return JSONResponse({"reply": "Unauthorized"}, status_code=403)
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        return JSONResponse({"reply": "Exam not found"}, status_code=404)
        
    # Gemini call
    model = _get_gemini_model()
    if not model:
        return JSONResponse({"reply": "Gemini AI is not configured. Please set GEMINI_API_KEY."})

    system_context = f"""You are an AI assistant for the Lexora IELTS platform.
The user is editing a mock exam titled '{exam.title}' (scope: {exam.test_scope}).
Help with generating questions, fixing extracted text, or IELTS advice. Be concise."""
    
    try:
        response = model.generate_content(system_context + "\nUser: " + chat_req.message)
        return JSONResponse({"reply": response.text})
    except Exception as e:
        return JSONResponse({"reply": f"Gemini error: {str(e)}"}, status_code=500)

@router.post("/{exam_id}/bulk-answers")
async def bulk_answers(request: Request, exam_id: int, bulk_text: str = Form(...), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Only Admins can save answers")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    import re
    # Matches "1. NOT GIVEN" or "1) NOT GIVEN" or "1 NOT GIVEN"
    matches = re.finditer(r'(\d+)[\.\)]?\s+([^\r\n]+)', bulk_text)
    
    count = 0
    
    for match in matches:
        try:
            q_num = int(match.group(1))
            ans_text = match.group(2).strip()
            
            # Find question
            q = db.query(Question).join(QuestionBlock).join(ExamSection).filter(
                ExamSection.exam_id == exam_id,
                Question.question_number == q_num
            ).first()
            
            if q:
                q.correct_answer_text = ans_text
                
                # If it's an MCQ, try to automatically set the correct AnswerOption
                if q.q_type == "MCQ":
                    db.query(AnswerOption).filter(AnswerOption.question_id == q.id).update({"is_correct": False})
                    # E.g., if ans_text is "C" or "C. Something", look for matching option
                    ans_letter = ans_text[0].upper() if ans_text else ""
                    
                    options = db.query(AnswerOption).filter(AnswerOption.question_id == q.id).all()
                    for opt in options:
                        if opt.text.strip().upper().startswith(ans_letter):
                            opt.is_correct = True
                            break
        except Exception:
            pass
            
    db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.post("/{exam_id}/update-section-content/{section_id}")
async def update_section_content(
    exam_id: int, 
    section_id: int,
    passage_text: str = Form(None),
    instructions: str = Form(None),
    db: SessionMaster = Depends(get_mdb)
):
    section = db.query(ExamSection).filter(ExamSection.id == section_id).first()
    if section:
        # For now, we assume one block per section for manual edits or update the first block
        if not section.blocks:
            from master_database import QuestionBlock
            block = QuestionBlock(section_id=section_id, passage_text=passage_text or "", instructions=instructions or "")
            db.add(block)
        else:
            block = section.blocks[0]
            if passage_text is not None: block.passage_text = passage_text
            if instructions is not None: block.instructions = instructions
        db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.post("/{exam_id}/delete-section/{section_id}")
async def delete_section(request: Request, exam_id: int, section_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    section = db.query(ExamSection).filter(ExamSection.id == section_id, ExamSection.exam_id == exam_id).first()
    if section:
        db.delete(section)
        db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

@router.post("/{exam_id}/delete-question/{question_id}")
async def delete_question(request: Request, exam_id: int, question_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["owner", "admin"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    question = db.query(Question).filter(Question.id == question_id).first()
    if question:
        db.delete(question)
        db.commit()
    return RedirectResponse(f"/mock/{exam_id}/build", status_code=303)

from services.scoring_engine import calculate_band, normalize_text
from datetime import datetime

@router.post("/{exam_id}/submit")
async def submit_exam(request: Request, exam_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=403, detail="Must be logged in to submit")
        
    exam = db.query(MockExam).filter(MockExam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    form = await request.form()
    
    # 1. Reuse existing in-progress attempt if possible
    attempt = db.query(MockAttempt).filter(
        MockAttempt.student_id == user.id,
        MockAttempt.exam_id == exam.id,
        MockAttempt.status == "in_progress"
    ).first()
    
    if not attempt:
        attempt = MockAttempt(
            tenant_id=user.tenant_id,
            student_id=user.id,
            exam_id=exam.id,
            status="completed",
            completed_at=datetime.utcnow()
        )
        db.add(attempt)
    else:
        attempt.status = "completed"
        attempt.completed_at = datetime.utcnow()
        
    db.flush()
    
    # 2. Grade
    raw_score = 0
    for key, val in form.items():
        if key.startswith("q"):
            # Could be GAP_FILL or MCQ
            q_id = int(key[1:])
            question = db.query(Question).filter(Question.id == q_id).first()
            if not question:
                continue
                
            is_correct = False
            option_id = None
            text_resp = ""
            
            if question.q_type == "MCQ":
                # val is the option_id chosen
                try:
                    option_id = int(val)
                    opt = db.query(AnswerOption).filter(AnswerOption.id == option_id).first()
                    if opt and opt.is_correct:
                        is_correct = True
                except ValueError:
                    pass
            else:
                # text answer
                text_resp = val
                expected = normalize_text(question.correct_answer_text)
                provided = normalize_text(text_resp)
                if expected and provided == expected:
                    is_correct = True
                    
            if is_correct:
                raw_score += question.points
                
            ans = db.query(AttemptAnswer).filter(
                AttemptAnswer.attempt_id == attempt.id,
                AttemptAnswer.question_id == q_id
            ).first()
            
            if not ans:
                ans = AttemptAnswer(attempt_id=attempt.id, question_id=q_id)
                db.add(ans)
                
            ans.text_response = text_resp
            ans.option_id = option_id
            ans.is_correct = is_correct

        elif key.startswith("writing_response_"):
            try:
                s_id = int(key.replace("writing_response_", ""))
                section = db.query(ExamSection).filter(ExamSection.id == s_id).first()
                if section and section.blocks and section.blocks[0].questions:
                    q_id = section.blocks[0].questions[0].id
                    ans = db.query(AttemptAnswer).filter(
                        AttemptAnswer.attempt_id == attempt.id,
                        AttemptAnswer.question_id == q_id
                    ).first()
                    if not ans:
                        ans = AttemptAnswer(attempt_id=attempt.id, question_id=q_id)
                        db.add(ans)
                    ans.text_response = val
            except Exception:
                pass
            
    # 3. Handle Writing Feedback (if any)
    feedback_type = form.get("feedback_type") # 'ai' or 'teacher'
    selected_teacher_id = form.get("teacher_id")
    
    if feedback_type:
        attempt.feedback_preference = feedback_type
        if selected_teacher_id:
            attempt.selected_teacher_id = int(selected_teacher_id)
            
        from master_database import WritingFeedback
        wf = WritingFeedback(
            attempt_id=attempt.id,
            feedback_type=feedback_type,
            teacher_id=int(selected_teacher_id) if selected_teacher_id else None,
            status="pending"
        )
        db.add(wf)
        
    # 4. Finalize Score
    attempt.total_score = raw_score
    attempt.band_score = calculate_band(raw_score, exam.test_scope)
    db.commit()
    
    return RedirectResponse(f"/mock/results/{attempt.id}", status_code=303)

@router.get("/results/{attempt_id}/export")
async def export_results(attempt_id: int, db: SessionMaster = Depends(get_mdb)):
    from docx import Document
    from docx.shared import Inches
    import io
    
    attempt = db.query(MockAttempt).filter(MockAttempt.id == attempt_id).first()
    if not attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")
        
    doc = Document()
    doc.add_heading(f"IELTS Exam Results: {attempt.exam.title}", 0)
    
    doc.add_paragraph(f"Candidate: {attempt.student.full_name or attempt.student.username}")
    doc.add_paragraph(f"Date: {attempt.completed_at.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Overall Band Score: {attempt.band_score}")
    
    doc.add_heading("Writing Responses", level=1)
    
    # Get writing responses
    for ans in attempt.answers:
        if ans.question.section_type and "Writing" in ans.question.section_type:
            doc.add_heading(f"{ans.question.section_type}", level=2)
            doc.add_paragraph(ans.text_response)
            doc.add_paragraph(f"Word Count: {len(ans.text_response.split())}")
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f"Lexora_Result_{attempt_id}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/results/{attempt_id}", response_class=HTMLResponse)
async def mock_results(request: Request, attempt_id: int, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=302)
        
    attempt = db.query(MockAttempt).filter(MockAttempt.id == attempt_id).first()
    if not attempt or (attempt.student_id != user.id and user.role != "owner"):
        raise HTTPException(status_code=404, detail="Result not found")
        
    return templates.TemplateResponse("mock_results.html", {
        "request": request,
        "user": user,
        "attempt": attempt,
        "active_page": "mock"
    })

@router.get("/history", response_class=HTMLResponse)
async def mock_history(request: Request, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=302)
        
    # Get all completed attempts for this user
    attempts = db.query(MockAttempt).filter(
        MockAttempt.student_id == user.id,
        MockAttempt.status == "completed"
    ).order_by(MockAttempt.completed_at.desc()).all()
    
    return templates.TemplateResponse("mock_history.html", {
        "request": request,
        "user": user,
        "attempts": attempts,
        "active_page": "history"
    })

@router.get("/inquiries", response_class=HTMLResponse)
async def writing_inquiries(request: Request, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=302)
    if user.role != "owner" and user.account_type != "teacher":
        raise HTTPException(status_code=403, detail="Teachers and Owner only")

    # Owner sees all, teacher sees only their own
    q = db.query(WritingFeedback)
    if user.role != "owner":
        q = q.filter(WritingFeedback.teacher_id == user.id)

    pending = q.filter(WritingFeedback.status == "pending").all()
    completed = q.filter(WritingFeedback.status == "completed").order_by(WritingFeedback.id.desc()).limit(20).all()

    def enrich(items):
        result = []
        for fb in items:
            attempt = db.query(MockAttempt).filter(MockAttempt.id == fb.attempt_id).first()
            if not attempt:
                continue
            student = db.query(User).filter(User.id == attempt.student_id).first()
            exam = db.query(MockExam).filter(MockExam.id == attempt.exam_id).first()
            writing_ans = db.query(AttemptAnswer).join(Question).join(QuestionBlock).join(ExamSection).filter(
                AttemptAnswer.attempt_id == attempt.id,
                ExamSection.section_type.ilike("%writing%")
            ).first()
            result.append({
                "feedback": fb, "student": student, "exam": exam,
                "attempt": attempt,
                "writing_text": writing_ans.text_response if writing_ans else "",
            })
        return result

    return templates.TemplateResponse("writing_inquiries.html", {
        "request": request,
        "user": user,
        "active_page": "writing_inquiries",
        "pending": enrich(pending),
        "completed": enrich(completed),
    })

@router.post("/inquiries/{feedback_id}/grade")
async def grade_writing(request: Request, feedback_id: int, score: float = Form(...), feedback_text: str = Form(...), db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or (user.role != "owner" and user.account_type != "teacher"):
        raise HTTPException(status_code=403, detail="Teachers only")

    fb = db.query(WritingFeedback).filter(WritingFeedback.id == feedback_id).first()
    if not fb:
        raise HTTPException(status_code=404)

    fb.score = score
    fb.feedback_text = feedback_text
    fb.status = "completed"
    db.commit()

    return RedirectResponse("/mock/inquiries", status_code=303)

@router.get("/classes/join", response_class=HTMLResponse)
async def join_class_page(request: Request):
    user = get_current_user(request)
    return templates.TemplateResponse("student_join_class.html", {
        "request": request,
        "user": user,
        "active_page": "my_classes"
    })
# --- TEACHER CLASS MANAGEMENT ---

@router.get("/classes", response_class=HTMLResponse)
async def unified_classes(request: Request, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user: return RedirectResponse("/login", status_code=302)
    
    if user.account_type == "teacher" or user.role == "owner":
        q = db.query(ClassRoom)
        if user.role != "owner":
            q = q.filter(ClassRoom.teacher_id == user.id)
        classes = q.all()
        return templates.TemplateResponse("teacher_classes.html", {
            "request": request, 
            "user": user, 
            "classes": classes, 
            "active_page": "classes"
        })
    else:
        # Student view: classes they are members of
        memberships = db.query(ClassMember).filter(ClassMember.student_id == user.id).all()
        classes = [m.classroom for m in memberships]
        return templates.TemplateResponse("student_classes.html", {
            "request": request, 
            "user": user, 
            "classes": classes, 
            "active_page": "classes"
        })

@router.get("/teacher/classes", response_class=HTMLResponse)
async def teacher_classes(request: Request, db: SessionMaster = Depends(get_mdb)):
    user = get_current_user(request)
    if not user or user.role not in ["teacher", "owner"]:
        return RedirectResponse("/login", status_code=302)
        
    classes = db.query(ClassRoom).filter(ClassRoom.teacher_id == user.id).all()
    return templates.TemplateResponse("teacher_classes.html", {
        "request": request,
        "user": user,
        "classes": classes,
        "active_page": "classes"
    })

@router.post("/teacher/classes/create")
async def create_class(
    request: Request, 
    class_name: str = Form(...), 
    db: SessionMaster = Depends(get_mdb)
):
    user = get_current_user(request)
    if not user or user.role not in ["teacher", "owner"]:
        raise HTTPException(status_code=403)
        
    import random, string
    invite_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
    
    new_class = ClassRoom(
        teacher_id=user.id,
        name=class_name,
        invite_code=invite_code
    )
    db.add(new_class)
    db.commit()
    return RedirectResponse("/mock/teacher/classes", status_code=303)

@router.post("/student/join-class")
async def join_class(
    request: Request, 
    invite_code: str = Form(...), 
    db: SessionMaster = Depends(get_mdb)
):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=302)
        
    classroom = db.query(ClassRoom).filter(ClassRoom.invite_code == invite_code.upper()).first()
    if not classroom:
        return RedirectResponse("/mock/history?error=Invalid invite code", status_code=303)
        
    # Check if already a member
    existing = db.query(ClassMember).filter(
        ClassMember.class_id == classroom.id,
        ClassMember.student_id == user.id
    ).first()
    
    if not existing:
        member = ClassMember(class_id=classroom.id, student_id=user.id)
        db.add(member)
        db.commit()
        
    return RedirectResponse("/mock/history?joined=success", status_code=303)


# ─── Writing / Speaking Inquiries (Teacher + Owner) ──────────────────────────

